import sys
import os
import json
import shutil
import csv
from dotenv import load_dotenv
from docx import Document
from docx.api import Document

from classes.rule_set import RuleSet
from classes.rule_set_legacy import RuleSetLegacy
import classes.globals as g

class RooDocument(object):
    def __init__(self):
        self.get_environment()
        self.get_chapter_codes()
        self.get_arguments()
        self.open_document()
        self.get_document_type()
        self.read_table()
        self.process_table()
        if self.modern:
            self.process_subdivisions()
        else:
            self.remove_working_nodes()
        self.write_table()
        self.kill_document()

    def get_environment(self):
        load_dotenv('.env')
        self.source_folder = os.getenv('source_folder')
        self.code_list = os.getenv('code_list')

    def get_chapter_codes(self):
        g.all_headings = {}
        g.all_subheadings = {}
        with open(self.code_list) as csv_file:
            csv_reader = csv.DictReader(csv_file)
            for row in csv_reader:
                if row["Commodity code"][-6:] == "000000" and row["Commodity code"][-8:] != "00000000":
                    heading = row["Commodity code"][0:4]
                    g.all_headings[heading] = row["Description"]

                if row["Commodity code"][-4:] == "0000" and row["Commodity code"][-6:] != "000000":
                    subheading = row["Commodity code"][0:6]
                    g.all_subheadings[subheading] = row["Description"]

        csv_file.close()

    def get_arguments(self):
        if len(sys.argv) > 1:
            self.docx_filename = sys.argv[1]
            self.docx_filepath = os.path.join(self.source_folder, self.docx_filename)
        else:
            print("Please supply an input document")
            sys.exit()

        self.export_folder = os.path.join(os.getcwd(), "export")
        self.export_filename = self.docx_filename.replace(".docx", "").replace(" ", "_").lower()
        self.export_filepath = os.path.join(self.export_folder, self.export_filename) + ".json"

    def open_document(self):
        self.document = Document(self.docx_filepath)

    def get_document_type(self):
        filename_compare = self.docx_filename.replace(" PSR.docx", "")
        modern_documents = [
            "EU", "Japan", "Turkey", "Canada"
        ]
        if filename_compare in modern_documents:
            self.modern = True
        else:
            self.modern = False

    def read_table(self):
        table = self.document.tables[0]
        if self.modern:
            rename_keys = {
                "Classification": "original_heading",
                "PSR": "original_rule"
            }
        else:
            rename_keys = {
                "Classification": "original_heading",
                "Description": "description",
                "PSR": "original_rule",
                "PSR2": "original_rule2"
            }

        data = []

        keys = None
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                keys = tuple(text)
                continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            data.append(row_data)

        self.rows = []
        for item in data:
            if item != {}:
                for old_key in rename_keys:
                    new_key = rename_keys[old_key]
                    item[new_key] = item.pop(old_key)

                self.rows.append(item)

    def process_table(self):
        if self.modern:
            self.process_table_modern()
        else:
            self.process_table_legacy()

    def process_table_modern(self):
        self.rule_sets = []
        for row in self.rows:
            if "Section" in row["original_heading"]:
                pass
            elif "SECTION" in row["original_heading"]:
                pass
            elif "Chapter" in row["original_heading"]:
                pass
            else:
                rule_set = RuleSet(row)
                self.rule_sets.append(rule_set.as_dict())

    def process_table_legacy(self):
        self.rule_sets = []
        for row in self.rows:
            rule_set = RuleSetLegacy(row)
            self.rule_sets.append(rule_set.as_dict())

        self.normalise_chapter()

    def normalise_chapter(self):
        for chapter in range(1, 98):
            if chapter == 20:
                a = 1
            is_chapter_mentioned = False
            rule_set_count = 0
            if chapter != 77:
                for rule_set in self.rule_sets:
                    if rule_set["chapter"] == chapter:
                        rule_set_count += 1
                        if rule_set["is_chapter"]:
                            is_chapter_mentioned = True

            if rule_set_count > 1 and is_chapter_mentioned:
                # Check if there are any ex codes anywhere except for on the chapter itself
                # If there are not, then we can just replicate rules, ignoring the headings that are exceptions
                has_ex_code_headings = False
                for rule_set in self.rule_sets:
                    if rule_set["chapter"] == chapter:
                        if "chapter" not in rule_set["heading"].lower():
                            if "ex" in rule_set["heading"]:
                                has_ex_code_headings = True
                                break

                if not has_ex_code_headings:
                    self.normalise_standard_chapter(chapter)
                else:
                    self.normalise_complex_chapter(chapter)

    def normalise_complex_chapter(self, chapter):
        # Get all the other rules in that chapter that are not the chapter heading
        matches = {}
        contains_subheading = False
        index = 0
        chapter_index = None
        if chapter == 20:
            a = 1
        for rule_set in self.rule_sets:
            if rule_set["chapter"] == chapter:
                if rule_set["is_chapter"]:
                    chapter_index = index
                    
                if rule_set["is_subheading"]:
                    contains_subheading = True

                matches[index] = rule_set
            index += 1

        # Loop through all of the headings in chapter (according to the DB)
        # If the heading is missing:
        #   add in the heading as a copy of the chapter rule_set
        
        chapter_string = str(chapter).rjust(2, "0")
        if contains_subheading == False:
            for heading in g.all_headings:
                if heading[0:2] == chapter_string:
                    heading_exists_in_rule_set = False
                    for match in matches:
                        if heading in matches[match]["headings"]:
                            heading_exists_in_rule_set = True
                            if matches[match]["is_ex_code"]:
                                process_ex_code = True
                            else:
                                process_ex_code = False
                            break
                        
                    if not heading_exists_in_rule_set:
                        # In expanding out the definition of a chapter to its subheadings, we will need
                        # to copy the rules from the chapter down to the headings that did not previously
                        # exist. In such cases, we need to create those headings and add them to the rule set.
                        obj = {
                            "heading": heading,
                            "headings": [heading],
                            "subheadings": [],
                            "chapter": chapter,
                            "subdivision": matches[chapter_index]["subdivision"],
                            "prefix": "",
                            "min": g.format_parts(heading, 0),
                            "max": g.format_parts(heading, 1),
                            "rules": matches[chapter_index]["rules"],
                            "is_ex_code": False,
                            "is_chapter": False,
                            "is_heading": True,
                            "is_subheading": False,
                            "is_range": False,
                            "valid": True
                        }
                        self.rule_sets.append(obj)
                    else:
                        if process_ex_code:
                            # This takes the definition of the chapter, reduced to 'Any other product'
                            # and assigns it to the matched heading as a counterpoint to the existing
                            # ex code, to cover all commodities that are not catered for within the
                            # specific ex code.
                            obj = {
                                "heading": matches[match]["heading"],
                                "headings": [],
                                "subheadings": [],
                                "chapter": chapter,
                                "subdivision": "Any other product",
                                "prefix": "",
                                "min": matches[match]["min"],
                                "max": matches[match]["max"],
                                "rules": matches[chapter_index]["rules"],
                                "is_ex_code": True,
                                "is_chapter": False,
                                "is_heading": True,
                                "is_subheading": False,
                                "is_range": False,
                                "valid": True
                            }
                            self.rule_sets.append(obj)
                            
                            # self.rule_sets[match]["subdivision"] = matches[match]["subdivision"]
                        
                    
                if heading[0:2] > chapter_string:
                    break
        else:
            a = 1
            # Firstly check for any headings under the chapter that have rules that are based on subheadings and not headings
            this_chapter_headings = []
            
            for heading in g.all_headings:
                heading_contains_matches = False
                if heading[0:2] == chapter_string:
                    subheadings = []
                    contains_subheadings = False
                    matched = False
                    for match in matches:
                        if heading in matches[match]["headings"]:
                            if matches[match]["is_subheading"]:
                                contains_subheadings = True
                                subheadings.append(matches[match])
                            matched = True

                    if matched:
                        status = "matched"
                    else:
                        status = "unmatched"
                        contains_subheadings = False
                    
                    obj = {heading: {
                        "status": status,
                        "contains_subheadings": contains_subheadings,
                        "subheadings": subheadings
                    }}
                    this_chapter_headings.append (obj)

                            
            a = 1
            
            for subheading in g.all_subheadings:
                if subheading[0:2] == chapter_string:
                    a = 1

        # Finally, remove the old chapter code, as its value has been copied elsewhere now
        self.rule_sets.pop(chapter_index)

    def normalise_standard_chapter(self, chapter):
        # First, get the chapter's own rule-set
        chapter_found = False
        string_to_find = "chapter " + str(chapter)
        index = -1
        for rule_set in self.rule_sets:
            index += 1
            if string_to_find in rule_set["heading"].lower():
                chapter_found = True
                rule_set_rubric = RuleSetLegacy()
                rule_set_rubric.rules = rule_set["rules"]
                break

        # Then, get all headings that do not have rulesets
        # and assign the rulesets to them
        chapter_string = str(chapter).rjust(2, "0")
        for heading in g.all_headings:
            if heading[0:2] == chapter_string:
                matched = False
                for rule_set in self.rule_sets:
                    if rule_set["heading"] == heading:
                        matched = True
                        break
                if not matched:
                    obj = {
                        "heading": heading,
                        "description": g.all_headings[heading],
                        "subdivision": "",
                        "is_ex_code": False,
                        "prefix": "",
                        "min": heading + "000000",
                        "max": heading + "999999",
                        "rules": rule_set["rules"],
                        "valid": True,
                        "chapter": chapter
                    }
                    self.rule_sets.append(obj)

            if heading[0:2] > chapter_string:
                break

        if chapter_found:
            self.rule_sets.pop(index)

    def process_subdivisions(self):
        previous_ruleset = None
        for rule_set in self.rule_sets:
            if rule_set["subdivision"] != "":
                if previous_ruleset is not None:
                    rule_set["heading"] = previous_ruleset["heading"]
                    rule_set["min"] = previous_ruleset["min"]
                    rule_set["max"] = previous_ruleset["max"]
            previous_ruleset = rule_set

    def write_table(self):
        out_file = open(self.export_filepath, "w")
        json.dump({"rule_sets": self.rule_sets}, out_file, indent=6)
        out_file.close()

        dest = "/Users/mattlavis/sites and projects/1. Online Tariff/ott prototype/app/data/roo/uk/psr_new/" + self.export_filename + ".json"
        shutil.copy(self.export_filepath, dest)

    def kill_document(self):
        self.document = None

    def remove_working_nodes(self):
        for rule_set in self.rule_sets:
            try:
                del rule_set["description"]
            except:
                pass

            try:
                del rule_set["headings"]
            except:
                pass

            try:
                del rule_set["subheadings"]
            except:
                pass

            try:
                del rule_set["is_ex_code"]
            except:
                pass

            try:
                del rule_set["is_chapter"]
            except:
                pass

            try:
                del rule_set["is_heading"]
            except:
                pass

            try:
                del rule_set["is_subheading"]
            except:
                pass

            try:
                del rule_set["is_range"]
            except:
                pass
